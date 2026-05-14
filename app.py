import streamlit as st
import fitz  # PyMuPDF
fitz.TOOLS.mupdf_display_errors(False)
import google.generativeai as genai
import io
import json
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =====================================================================
# 1. CONFIGURACIÓN E INTERFAZ (UI)
# =====================================================================
st.set_page_config(page_title="BioCardio Clinical Assistant", page_icon="🏥", layout="wide")

st.markdown("""
    <style>
    .stApp { background-color: #f4f7f4; }
    .stFileUploader label { font-weight: bold !important; color: #1a1a1a !important; font-size: 19px !important; display: block; margin-bottom: 12px !important; }
    .step-container { background-color: white; padding: 25px 30px; border-radius: 15px; border-left: 8px solid #007d32; box-shadow: 0 4px 12px rgba(0,0,0,0.08); margin-bottom: 35px; }
    .step-header { color: #007d32; font-size: 24px; font-weight: bold; margin-top: 15px; margin-bottom: 10px; }
    .stButton>button { background-color: #007d32 !important; color: white !important; border-radius: 12px !important; height: 55px !important; font-size: 18px !important; font-weight: bold !important; width: 100%; transition: 0.3s; border: none !important; }
    </style>
""", unsafe_allow_html=True)

# =====================================================================
# 2. MOTOR DE EXTRACCIÓN DE TABLAS MATRICIALES 
# =====================================================================
def extraer_texto_matricial(pdf_bytes, paginas_str=""):
    """ Extrae las tablas manteniendo las columnas intactas mediante barras (|) """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    texto = ""
    try:
        p_list = []
        if paginas_str.strip():
            for b in paginas_str.split(','):
                if '-' in b:
                    s, e = b.split('-')
                    p_list.extend(range(int(s), int(e)+1))
                else: 
                    p_list.append(int(b))
        else: 
            p_list = range(1, len(doc) + 1)
        
        for p in p_list:
            if 0 < p <= len(doc):
                page = doc[p-1]
                texto += f"\n--- PÁGINA {p} ---\n"
                tablas = page.find_tables()
                if tablas and len(tablas.tables) > 0:
                    for tabla in tablas.tables:
                        matriz = tabla.extract()
                        for fila in matriz:
                            fila_limpia = [" ".join(str(c).split()) if c else "VACIO" for c in fila]
                            texto += " | ".join(fila_limpia) + "\n"
                else:
                    texto += page.get_text("text") + "\n"
    except Exception as e: 
        texto = f"Error lectura: {e}"
    return texto

# =====================================================================
# 3. GENERADOR DEL DOCUMENTO WORD (PLANTILLAS ESTRICTAS)
# =====================================================================
def crear_documento_word(datos_json, protocolo_nombre):
    try:
        match = re.search(r'\{.*\}', datos_json, re.DOTALL)
        datos = json.loads(match.group(0)) if match else json.loads(datos_json)
    except:
        datos = {"visita": "Error", "procedimientos": {}, "otros_procedimientos": [], "detalles": {}}

    proc = datos.get("procedimientos", {})
    otros = datos.get("otros_procedimientos", [])
    det = datos.get("detalles", {})
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0, 0, 0)
    for i in range(1, 4):
        if f'Heading {i}' in doc.styles:
            h_style = doc.styles[f'Heading {i}']
            h_style.font.name = 'Calibri'
            h_style.font.bold = True
            h_style.font.color.rgb = RGBColor(0, 0, 0)

    es_alexion = "ALXN" in protocolo_nombre

    # ---------------------------------------------------------
    # PLANTILLA A: ALEXION (ALXN2220-ATTR-CM-301)
    # ---------------------------------------------------------
    if es_alexion:
        doc.add_heading(f"HOJA DE ENFERMERÍA: {datos.get('visita', 'N/A')}", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Protocolo: {protocolo_nombre}     ID: ______________").bold = True
        
        doc.add_heading("PRIMEROS PASOS", level=2)
        doc.add_paragraph("☐ Registrar visita para asignar infusión al paciente en Almac (IXRS)")
        doc.add_paragraph("☐ Fecha de la visita: ______________")
        
        doc.add_heading("PRE-INFUSIÓN (dentro de las 2 horas antes de la infusión)", level=2)
        
        if proc.get("signos_vitales"):
            doc.add_paragraph("☐ Peso: _________ kg")
            
        if proc.get("test_embarazo"):
            doc.add_paragraph("☐ Test de embarazo en orina (mujeres en edad fértil)").bold = True

        if proc.get("cuestionarios"):
            doc.add_paragraph("☐ Cuestionarios o PROs (TABLET en armario):").bold = True
            doc.add_paragraph("    ☐ KCCQ-OS  ☐ EQ-5D-5L  ☐ SF-36  ☐ PGIC (EN PAPEL)")

        hdr = ["Tiempo", "Hora", "PA Sist.", "PA Diast.", "Pulso", "Resp.", "O2 (%)", "Ta (ºC)"]
        if proc.get("signos_vitales"):
            doc.add_paragraph("☐ Signos vitales: (tras 5 minutos descanso)").bold = True
            rows = 3 if proc.get("laboratorio") else 2
            t_vit_pre = doc.add_table(rows=rows, cols=8)
            t_vit_pre.style = 'Table Grid'
            for i, h in enumerate(hdr): t_vit_pre.cell(0, i).text = h
            if proc.get("laboratorio"):
                t_vit_pre.cell(1, 0).text = "Pre-extracción"
                t_vit_pre.cell(2, 0).text = "Pre-infusión"
            else:
                t_vit_pre.cell(1, 0).text = "Pre-infusión"
            doc.add_paragraph("")

        if proc.get("ecg_pre"):
            doc.add_paragraph("☐ Pre-Electrocardiograma de 12 derivaciones").bold = True
            doc.add_paragraph("    ☐ Posición supino    ☐ FC: ____  ☐ PR: ____  ☐ QRS: ____  ☐ QT: ____  ☐ QTc: ____")

        if proc.get("eco"):
            doc.add_paragraph("☐ Ecocardiograma (dejar instrucciones) → realizado por ____________").bold = True

        if proc.get("laboratorio"):
            doc.add_paragraph("☐ Extraer muestras de sangre y orina").bold = True
            doc.add_paragraph(f"Tubos a extraer y preparar: {det.get('laboratorio_tubos', 'Ver manual')}").italic = True

        if proc.get("test_6mwt"):
            doc.add_paragraph("☐ Test de los 6 minutos (6MWT)").bold = True
            doc.add_paragraph("    • Rellenar en Course Name: 6MWT 2F; y en Course Lenght: 20m")

        if proc.get("infusion"):
            doc.add_heading("ADMINISTRACIÓN DE INFUSIÓN + SIGNOS VITALES", level=2)
            doc.add_paragraph("*La infusión debe durar 1 hora. Limpiar vías con dextrosa 5%.").italic = True
            doc.add_paragraph("☐ HORA DE INICIO: _________   ☐ HORA DE FIN: _________")
            doc.add_paragraph("\nSignos vitales durante la infusión:").bold = True
            t_vit_inf = doc.add_table(rows=5, cols=8)
            t_vit_inf.style = 'Table Grid'
            for i, h in enumerate(hdr): t_vit_inf.cell(0, i).text = h
            for r, text in enumerate(["15 min infusión", "30 min infusión", "45 min infusión", "1h infusión"]):
                t_vit_inf.cell(r+1, 0).text = text
            doc.add_paragraph("")

            doc.add_heading("POST-INFUSIÓN (30 min)", level=2)
            if proc.get("pk_post"): doc.add_paragraph("☐ PK-post infusión: sacar del brazo opuesto.").bold = True
            if proc.get("ecg_post"): doc.add_paragraph("☐ Post-Electrocardiograma.").bold = True
            
            t_vit_post = doc.add_table(rows=2, cols=8)
            t_vit_post.style = 'Table Grid'
            for i, h in enumerate(hdr): t_vit_post.cell(0, i).text = h
            t_vit_post.cell(1, 0).text = "30 min post"
            doc.add_paragraph("")
        
        doc.add_paragraph("\nFirmado: _______________________      Fecha: ______________")

        # --- PÁGINA 2: MÉDICO ---
        doc.add_page_break()
        doc.add_heading(f"HOJA MÉDICA: {datos.get('visita', 'N/A')}", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Protocolo: {protocolo_nombre}     ID: ______________").bold = True
        doc.add_paragraph("☐ Fecha de la visita: ______________\n")
        
        if proc.get("examen_fisico"): 
            doc.add_paragraph("☐ Examen físico COMPLETO").bold = True
            doc.add_paragraph("(Ver aspecto general, piel, nariz, orejas, ojos, cuello, garganta, corazón, abdomen, pulmones, sistema vascular, sistema nervioso, sistema musculo esquelético y extremidades)\n")
        
        if proc.get("infusion") or proc.get("examen_fisico_breve"): 
            doc.add_paragraph("☐ Examen físico BREVE post-infusión (Dirigido por síntomas)").bold = True
            doc.add_paragraph("Inclusive of general appearance, heart, lungs, skin, musculoskeletal system and extremities and other organs or body systems as clinically indicated should be performed prior to the participant's discharge.\n")

        if proc.get("nyha"): 
            doc.add_paragraph("☐ Clasificación funcional NYHA (seleccionar una):").bold = True
            doc.add_paragraph("   ☐ NYHA I: Asintomático")
            doc.add_paragraph("   ☐ NYHA II: Falta de aire (disnea) a grandes esfuerzos")
            doc.add_paragraph("   ☐ NYHA III: Falta de aire (disnea) a pequeños esfuerzos")
            doc.add_paragraph("   ☐ NYHA IV: Falta de aire (disnea) en reposo (se ahoga estando quieto)\n")

        if proc.get("karnofsky"): 
            doc.add_paragraph("☐ Discapacidad (con escala Karnofsky):").bold = True
            doc.add_paragraph("   100% - Actividad normal (capaz de desempeñar actividades), asintomático.")
            doc.add_paragraph("   90% - Actividad normal, con síntomas y signos leves.")
            doc.add_paragraph("   80% - Actividad normal con esfuerzo, síntomas leves.")
            doc.add_paragraph("   70% - Capaz de cuidar de si mismo, pero no realiza trabajo activo.")
            doc.add_paragraph("   60% - En ocasiones necesita ayuda, capaz de cuidarse la mayor parte del tiempo.")
            doc.add_paragraph("   50% - Necesita atención médica y ayuda frecuente.")
            doc.add_paragraph("   40% - Con discapacidad, requiere cuidados especiales.")
            doc.add_paragraph("   30% - Discapacidad grave, en condiciones de hospitalización.")
            doc.add_paragraph("   20% - Enfermo grave, necesita tratamiento activo de sostén.")
            doc.add_paragraph("   10% - Paciente decaído o moribundo.")
            doc.add_paragraph("   0% - Paciente fallecido.\n")
        
        doc.add_paragraph("☐ Medicamentos, terapias o procedimientos simultáneos").bold = True
        doc.add_paragraph("☐ Eventos adversos (AEs)\n").bold = True
        doc.add_paragraph("\nFirmado (Médico): _______________________      Fecha: ______________")

# ---------------------------------------------------------
    # PLANTILLA B: ALNYLAM (ALN-TTRSC04-003)
    # ---------------------------------------------------------
    else:
        doc.add_heading(f"{datos.get('visita', 'N/A')} - Hoja de Enfermería", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("______________________________________________________________")
        doc.add_paragraph(f"Protocol: {protocolo_nombre}\nID:\nFecha de la visita:").bold = True
        doc.add_paragraph("______________________________________________________________\n")
        
        if proc.get("signos_vitales"):
            doc.add_paragraph("☐ Signos vitales:").bold = True
            hdr = ["Tiempo", "Hora", "PA Sist.", "PA Diast.", "Pulso", "Resp.", "O2 (%)", "Ta (ºC)"]
            t_vit = doc.add_table(rows=2, cols=8)
            t_vit.style = 'Table Grid'
            for i, h in enumerate(hdr): t_vit.cell(0, i).text = h
            doc.add_paragraph("")

        if proc.get("ecg_pre"):
            doc.add_paragraph("☐ Electrocardiograma de 12 derivaciones").bold = True
            doc.add_paragraph("    ☐ FC: ____  ☐ PR: ____  ☐ QRS: ____  ☐ QT: ____  ☐ QTc: ____")

        if proc.get("eco"): doc.add_paragraph("☐ Ecocardiograma").bold = True
        if proc.get("cuestionarios"): doc.add_paragraph("☐ Cuestionarios y PROs realizados.").bold = True
        if proc.get("test_embarazo"): doc.add_paragraph("☐ Test de embarazo en orina").bold = True

        if proc.get("laboratorio"):
            doc.add_paragraph("☐ Extraer muestras de sangre y orina para analítica central").bold = True
            doc.add_paragraph(f"{det.get('laboratorio_tubos', 'Ver manual de laboratorio')}").italic = True
            doc.add_paragraph("\n☐ Procesamiento de muestras en el laboratorio (según manual)").bold = True
            doc.add_paragraph("☐ Envío de muestras a Tª ambiente y congeladas").bold = True

        if proc.get("infusion"):
            doc.add_paragraph("\n☐ Administración de medicación del estudio (Study drug administration)").bold = True

        if otros:
            doc.add_paragraph("\n☐ OTROS PROCEDIMIENTOS PROGRAMADOS:").bold = True
            for proc_extra in otros:
                doc.add_paragraph(f"    ☐ {proc_extra}")

        doc.add_paragraph("\n\nFirma: ______________________________        Fecha: ______________")

        # --- PÁGINA 2: MÉDICO ---
        doc.add_page_break()
        doc.add_heading(f"{datos.get('visita', 'N/A')} - General / Médica", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("______________________________________________________________")
        doc.add_paragraph(f"Protocol: {protocolo_nombre}\nID:\nFecha de la visita:").bold = True
        doc.add_paragraph("______________________________________________________________\n")
        
        doc.add_paragraph("☐ Dar cita para próxima visita: ___________________").bold = True
        doc.add_paragraph("\n☐ Medicamentos concomitantes, terapias o procedimientos actuales:").bold = True
        doc.add_paragraph("\n☐ Eventos adversos desde la firma del consentimiento:").bold = True

        # Exámenes físicos completamente separados
        if proc.get("examen_fisico"): 
            doc.add_paragraph("\n☐ Examen físico COMPLETO (Full physical examination)").bold = True
        if proc.get("examen_fisico_breve"): 
            doc.add_paragraph("\n☐ Examen físico BREVE (Symptom-directed physical examination)").bold = True

        if proc.get("nyha"): 
            doc.add_paragraph("☐ Clasificación funcional NYHA evaluada:").bold = True
            doc.add_paragraph("   ☐ NYHA I  ☐ NYHA II  ☐ NYHA III  ☐ NYHA IV\n")

        if proc.get("karnofsky"): 
            doc.add_paragraph("☐ Puntuación Karnofsky (Karnofsky Performance Status):").bold = True
            doc.add_paragraph("   ☐ 100% ☐ 90% ☐ 80% ☐ 70% ☐ 60% ☐ 50% ☐ 40% ☐ 30% ☐ 20% ☐ 10% ☐ 0%\n")

        if proc.get("test_6mwt"): doc.add_paragraph("☐ Test de los 6 minutos (6MWT) completado").bold = True

        doc.add_paragraph("\nCIERRE DE VISITA:").bold = True
        doc.add_paragraph("☐ Registrar próxima visita en Greenphire")
        doc.add_paragraph("☐ Escribir historia clínica en Diraya (imprimir, fechar y firmar)")
        doc.add_paragraph("☐ Rellenar CRF (Medidata)")
        
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# =====================================================================
# 4. APLICACIÓN STREAMLIT (LÓGICA PRINCIPAL)
# =====================================================================
col_logo, col_titulo = st.columns([1, 5])
with col_logo:
    try: st.image("huj.png", width=140)
    except: st.title("🏥")
with col_titulo:
    st.markdown("<h1 style='color: #007d32; margin-top: 10px;'>BioCardio Clinical Assistant</h1>", unsafe_allow_html=True)
    st.markdown("<p style='font-size: 18px; color: #555;'>Herramienta profesional de gestión de ensayos clínicos</p>", unsafe_allow_html=True)

with st.sidebar:
    st.markdown("### 🔑 Configuración")
    st.markdown("Para que el sistema funcione, necesitas una **API Key** personal.")
    st.markdown("[Consigue tu clave gratuita en Google AI Studio](https://aistudio.google.com/app/apikey)")
    api_key = st.text_input("Introduce API Key:", type="password")
    modelo_seleccionado = st.selectbox("Modelo:", ["gemini-2.5-flash", "gemini-2.0-flash-lite"])

st.markdown('<div class="step-header">📄 Paso 1: Documentación</div>', unsafe_allow_html=True)
with st.container():
    st.markdown('<div class="step-container">Sube el protocolo (SoE) y los manuales de laboratorio.</div>', unsafe_allow_html=True)
    f_proto = st.file_uploader("1. SUBIR PROTOCOLO (SoE)", type=["pdf"])
    f_labs = st.file_uploader("2. SUBIR MANUALES LABORATORIO", type=["pdf"], accept_multiple_files=True)

st.markdown('<div class="step-header">🔍 Paso 2: Configuración de la Visita</div>', unsafe_allow_html=True)
with st.container():
    protocolo_sel = st.selectbox("Protocolo a maquetar:", ["Alexion (ALXN2220-ATTR-CM-301)", "Alnylam (ALN-TTRSC04-003)"])
    c1, c2 = st.columns(2)
    p_tabla = c1.text_input("Páginas Tabla SoE (ej. 23-28):")
    v_proto = c1.text_input("Visita Protocolo (ej. V 30/W108):")
    p_ass = c2.text_input("Páginas Assessments (ej. 67-80):")
    v_lab = c2.text_input("Visita Manual Lab:", "Visit 30/W108")

st.markdown('<div class="step-header">📋 Paso 3: Generación del Checklist</div>', unsafe_allow_html=True)
with st.container():
    if st.button("GENERAR DOCUMENTO DE VISITA"):
        if not api_key or not f_proto: 
            st.error("⚠️ Faltan documentos o la API Key.")
        else:
            barra_progreso = st.progress(0)
            status_text = st.empty()
            
            try:
                status_text.text("Analizando celdas exactas de la tabla SoE...")
                p_bytes = f_proto.read()
                t_soe = extraer_texto_matricial(p_bytes, p_tabla)
                t_ass = extraer_texto_matricial(p_bytes, p_ass)
                barra_progreso.progress(40)
                
                status_text.text("Leyendo el manual de laboratorio...")
                t_lab = ""
                if f_labs:
                    for f in f_labs: 
                        t_lab += extraer_texto_matricial(f.read(), "")
                barra_progreso.progress(70)
                
                status_text.text("IA aplicando razonamiento...")
                genai.configure(api_key=api_key)
                model = genai.GenerativeModel(modelo_seleccionado)
                
                prompt = f"""
                Eres un auditor clínico estricto. Tu trabajo es leer una TABLA SOE (matriz de datos) y extraer EXACTAMENTE los procedimientos marcados para la visita '{v_proto}'.
                
                INSTRUCCIONES DE RAZONAMIENTO OBLIGATORIO (Escribe tu razonamiento antes de generar el JSON):
                1. Busca la fila de encabezados y encuentra la columna exacta de '{v_proto}'.
                2. Recorre las filas hacia abajo mirando ÚNICAMENTE el contenido de esa columna.
                3. Si la celda en esa columna está escrita como "VACIO", el procedimiento es FALSE. Si tiene una 'X', asterisco o número, es TRUE.
                4. Haz una lista confirmando cada procedimiento.

                TABLA SOE MATRICIAL:
                {t_soe}
                
                MANUAL DE LABORATORIO:
                {t_lab[:60000]}
                
                REGLAS DE MAPEO (SOLO SI LA CELDA DE LA COLUMNA '{v_proto}' TIENE MARCA):
                - cuestionarios: true SOLO si hay marca en KCCQ, EQ-5D, SF-36, etc.
                - signos_vitales: true SOLO si hay marca en Vital signs o Weight.
                - ecg_pre: true SOLO si hay marca en 12-lead ECG.
                - ecg_post: true SOLO si el protocolo exige explícitamente ECG post-infusión ese día.
                - test_6mwt: true SOLO si hay marca en 6-Minute Walk Test.
                - laboratorio: true SOLO si hay marca en Central clinical laboratory tests. (Si es true, busca la visita '{v_lab}' en el texto del manual. PARA LOS COLORES DE LOS TUBOS USA OBLIGATORIAMENTE ESTA GUÍA EXCLUSIVA: Coagulación = Tapón azul (Na Citrate); Chemistry, cardiac biomarker, hep A = Tapón rojo (Serum tube clot activator); LFTs, Serum potassium, serum hcg, fsh, Vitamina A, ANA, ASMA, LKM 1, Anti-mitochondrial Ab, Anti-SLA = Tapón rojo (SST); HCV Confirmation = Tapón amarillo (SST); ADA, TTR protein, future research serum = Tapón rojo (Serum tube clot activator); Hematology, NfL, Genotyping, Plasma PK, future research plasma = Tapón morado (K2EDTA); Urinalysis, future research urine = Bote de orina. ESTÁ ESTRICTAMENTE PROHIBIDO dar explicaciones conversacionales como "El manual no detalla los colores". Limítate a cruzar la prueba requerida con esta guía y escribir el tubo correspondiente).
                - infusion: true SOLO si hay marca en Study intervention infusion.
                - pk_post: true SOLO si hay marca en PK samples.
                - eco: true SOLO si hay marca en Echocardiogram o ECHO.
                - examen_fisico: true SOLO si hay marca en 'Full physical examination'.
                - examen_fisico_breve: true SOLO si hay marca en 'Symptom-directed physical examination', 'Targeted physical examination' o 'Symptom-directed'.
                - test_embarazo: true SOLO si hay marca explícita en 'Pregnancy test' en la columna '{v_proto}'.
                - nyha: true SOLO si hay marca en NYHA.
                - karnofsky: true SOLO si hay marca en Karnofsky.

                ¡REGLA DE DOBLE VERIFICACIÓN Y EXTRACTOR UNIVERSAL (ESPECIAL ALNYLAM)!:
                Para CUALQUIER OTRA FILA de la tabla (ej. 'Vitamin A') que tenga una marca (X) en la columna '{v_proto}', captura el nombre exacto de ese procedimiento y añádelo a "otros_procedimientos".
                CUIDADO EXTREMO: Comprueba DOS VECES que la celda de esa prueba en la columna '{v_proto}' NO esté vacía. Si está "VACIO" o en blanco, PROHIBIDO incluirlo. No incluyas pruebas de visitas adyacentes.

                Tras tu razonamiento, genera el JSON estricto con este formato:
                {{
                  "visita": "{v_proto}",
                  "procedimientos": {{
                    "cuestionarios": false, "signos_vitales": false, "ecg_pre": false, "ecg_post": false, "laboratorio": false, "infusion": false,
                    "examen_fisico": false, "examen_fisico_breve": false, "test_embarazo": false,
                    "nyha": false, "karnofsky": false, "test_6mwt": false, "pk_post": false, "eco": false
                  }},
                  "otros_procedimientos": ["Nombre de prueba 1", "Nombre de prueba 2"],
                  "detalles": {{ "laboratorio_tubos": "Resumen directo extraído con colores..." }}
                }}
                """
                
                res = model.generate_content(prompt)
                
                doc_word = crear_documento_word(res.text, protocolo_sel)
                
                barra_progreso.progress(100)
                status_text.success("✅ Generado con éxito.")
                st.download_button("⬇️ Descargar Documento de Visita", doc_word, f"{v_proto}.docx")
                
            except Exception as e: 
                st.error(f"Error crítico en el proceso: {e}")
